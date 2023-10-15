import os
from docx import Document
from openpyxl import load_workbook
import psycopg2
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
config_path = os.path.expanduser("~/Desktop/config/config.json")


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def set_text_color_to_white(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_cell_margins(table, left=0, right=0):

    tc = table._element
    tblPr = tc.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    kwargs = {"left":left, "right":right}
    for m in ["left","right"]:
        node = OxmlElement("w:{}".format(m))
        node.set(qn('w:w'), str(kwargs.get(m)))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)

    tblPr.append(tblCellMar)

def set_table_width_to_page_width(table, page_width):
    table.width = page_width
def merge_empty_cells_in_first_row(table):
    # Get the first row
    first_row = table.rows[0]

    # Initialize variables to keep track of the merged cell range
    empty_col_idx = None
    non_empty_idx = None
    merged_text = []  # To store text from empty cells without introducing newlines

    for col_idx, cell in enumerate(first_row.cells):
        cell_value = cell.text.strip()  # Remove leading/trailing spaces
        if cell_value == "":
            if non_empty_idx is not None:
                empty_col_idx = col_idx
            if empty_col_idx is not None:
                merged_text.append(table.cell(0, empty_col_idx).text)  # Collect text from empty cells
                # Merge empty cells from start_col_idx to col_idx
                for merge_col_idx in range(non_empty_idx, empty_col_idx):
                    table.cell(0, merge_col_idx).merge(table.cell(0, empty_col_idx))
                empty_col_idx = None
        else:
            non_empty_idx = col_idx

    # Check if there are trailing empty cells to merge and append their text to the last non-empty cell
    if empty_col_idx is not None and non_empty_idx is not None:
        merged_text.append(table.cell(0, empty_col_idx).text)
        for merge_col_idx in range(non_empty_idx, empty_col_idx):
            table.cell(0, merge_col_idx).merge(table.cell(0, empty_col_idx))

# def set_text_alignment_left_with_padding(table):
#     for row_idx, row in enumerate(table.rows):
#         if row_idx == 0:
#             continue  # Skip the first row
#         for col_idx, cell in enumerate(row.cells):
#             if col_idx == 0:
                
#                 for paragraph in cell.paragraphs:
#                     paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def remove_newlines_from_first_row(table):
    first_row = table.rows[0]
    
    for col_idx, cell in enumerate(first_row.cells):
        text = cell.text.replace('\n', ' ').strip()
        cell.text = text

is_colored = False  # Переменная для отслеживания состояния окраски

def process_table(new_document, table_name, tables_directory , topic_color , cell_color):
    global is_colored
    excel_file_path = os.path.join(tables_directory, f"{table_name}.xlsx")
    print(table_name)
    try:
        # Load the Excel file
        excel_workbook = load_workbook(excel_file_path, data_only=True)
        excel_sheet = excel_workbook.active
        if excel_sheet.max_row > 0 and excel_sheet.max_column > 0:
            num_rows = excel_sheet.max_row
            num_cols = excel_sheet.max_column
            table = new_document.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for row_idx, row in enumerate(table.rows):
                is_colored = not is_colored  # Declare is_colored as a global variable

                for col_idx, cell in enumerate(row.cells):
                    if row_idx == 0:
                        
                        set_text_color_to_white(cell)
                    shading_color = topic_color if row_idx == 0 else cell_color if is_colored else "#fff"



                    cell.paragraphs[0].style.font.size = Pt(8)
                    cell.paragraphs[0].style.font.name = 'Arial' 

                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}" />')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
             
                    
            for row_idx, row in enumerate(excel_sheet.iter_rows()):
                
                for col_idx, cell in enumerate(row):
                    cell_value = cell.value if cell.value is not None else ""
                    # Check if cell_value is a numeric type and contains a period
                    if isinstance(cell_value, (float, int)) and '.' in str(cell_value):
                        cell_value = str(cell_value).replace('.', ',')  # Replace periods with commas

                    cell = table.cell(row_idx, col_idx)
                    cell.text = str(cell_value)
                    cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.width = Inches(1)  # Fixed width
                    cell.autofit = True  # Autofit the cell to its content

                    if row_idx == 0:
                        for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                            if paragraph.runs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 255, 255)
            new_document.add_paragraph()
            
            merge_empty_cells_in_first_row(table)
            remove_newlines_from_first_row(table)
            # set_text_alignment_left_with_padding(table)

            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if row_idx == 0:
                        
                        set_text_color_to_white(cell)
            set_cell_margins(table)
        else:
            new_document.add_paragraph("Empty table: " + table_name)
            new_document.add_paragraph()
    except FileNotFoundError:
        new_document.add_paragraph("Excel file not found: " + table_name)
        new_document.add_paragraph()
    except Exception as e:
        print(f"Error processing table {table_name}: {str(e)}")

def main():
# Check if the JSON config file existss
    conn = None
    if os.path.exists(config_path):
        with open(config_path, 'r') as config_file:
            config = json.load(config_file)
        # Check if the JSON config file contains the required information
            dbname = config['database']['name']
            user = config['database']['user']
            password = config['database']['password']



            # Set up a connection to the PostgreSQL database
            conn = psycopg2.connect(
                dbname=dbname,
                user=user,
                password=password,
                host="localhost",
                port="5432"
            )
    else:
        print("JSON config file not found at the specified path.")
    # Path to the directory where Excel files are stored
    tables_directory = "./tables_test"

    # Path to the output document
    output_document_path = "./word/BulletTestTemplate.docx"

    # Create a new Word document
    new_document = Document(output_document_path)

    # Create a cursor for executing SQL queries
    cursor = conn.cursor()

    # Fetch topics from the database
    cursor.execute("SELECT * FROM api_topic")
    topics = cursor.fetchall()

    for i, topic in enumerate(topics):
        topic_id, topic_name = topic
        topic_heading = f"{i + 1}. {topic_name}"
        topic_paragraph = new_document.add_paragraph()
        topic_run = topic_paragraph.add_run(topic_heading)
        topic_font = topic_run.font
        topic_font.name = 'Arial'
        topic_font.size = Pt(16)
        topic_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        topic_font.bold = True
        new_document.add_paragraph()

        cursor.execute("SELECT * FROM api_economic_index WHERE macro_topic_id = %s", (topic_id,))
        economic_indexes = cursor.fetchall()

        for k, economic_index in enumerate(economic_indexes):
            economic_index_id, economic_index_name, *some = economic_index
            economic_index_heading = f"{i + 1}.{k+1}.  {economic_index_name}"
            economic_index_paragraph = new_document.add_paragraph()
            economic_index_run = economic_index_paragraph.add_run(economic_index_heading)
            economic_index_font = economic_index_run.font
            economic_index_font.name = 'Arial'
            economic_index_font.size = Pt(12)
            economic_index_font.bold = True

            economic_index_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            economic_index_font.color.rgb = RGBColor(0, 0, 139)

            cursor.execute("SELECT path FROM api_table WHERE macro_economic_index_id = %s", (economic_index_id,))
            table_names = cursor.fetchall()
            
            topic_color ="#006FC0"
            cell_color = "#DDEBF7"
            if topic_name == "РЕГИОНЫ":
                topic_color = "#538235"
                cell_color = "#E1EEDA"
            elif topic_name == "ЗЕЛЕНАЯ ЭКОНОМИКА":
                topic_color = "#00B050"
                cell_color = "#EAF1DD"
            elif topic_name == "НАЦИОНАЛЬНЫЙ ПЛАН РАЗВИТИЯ РЕСПУБЛИКИ КАЗАХСТАН":
                topic_color = "#7030A0"
                cell_color = "#fff"
            for table_name in table_names:
                set_table_width_to_page_width(new_document.add_table(rows=1, cols=1), Inches(8.5))  # Set table width
                process_table(new_document, table_name[0], tables_directory , topic_color, cell_color)

    new_document.save("./word/test3.docx")

    cursor.close()
    conn.close()

if __name__ == "__main__":
    main()
