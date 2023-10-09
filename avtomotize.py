import os
from docx import Document
from openpyxl import load_workbook

# Set the path to the directory with Excel tables
directory_path = "./tables/"

# Get a list of all files in the specified directory
files = os.listdir(directory_path)

# Set the path to the Word document where you want to insert Excel tables
word_document_path = "./word/Bulletn2023_august.docx"
# Set the path for the new Word document
output_document_path = "./word/bulletinTest.docx"

# Check if the output document exists and delete it if it does
if os.path.exists(output_document_path):
    os.remove(output_document_path)

# Create a new Word document
new_document = Document()

# Initialize a variable to track whether the previous paragraph had anchor text
prev_paragraph_had_anchor = False

# Iterate through paragraphs to replace anchor text with tables
for paragraph in Document(word_document_path).paragraphs:
    for file_name in files:
        file_name_without_extension = os.path.splitext(file_name)[0]
        if paragraph.text.strip() == file_name_without_extension:
            print(f"Replacing paragraph text: '{paragraph.text}'")

            # Find the index of the matched file name
            file_index = files.index(file_name)

            # Get the path to the corresponding Excel file
            excel_file_path = os.path.join(directory_path, files[file_index])

            # Load the Excel file
            excel_workbook = load_workbook(excel_file_path, data_only=True)
            excel_sheet = excel_workbook.active

            # Create a new table with the same number of rows and columns as the Excel sheet
            num_rows = excel_sheet.max_row
            num_cols = excel_sheet.max_column
            table = new_document.add_table(rows=num_rows, cols=num_cols)

            # Populate the table with data from Excel
            for row_idx, row in enumerate(excel_sheet.iter_rows()):
                for col_idx, cell in enumerate(row):
                    table.cell(row_idx, col_idx).text = str(cell.value)

            # Add a new paragraph to separate the table from the previous content
            new_document.add_paragraph()

            # Set the variable to True to skip appending the paragraph text in the next iteration
            prev_paragraph_had_anchor = True

    if not prev_paragraph_had_anchor:
        # If the previous paragraph didn't have anchor text, append its text
        new_document.add_paragraph(paragraph.text)

    # Reset the variable for the next paragraph
    prev_paragraph_had_anchor = False

new_document.save(output_document_path)